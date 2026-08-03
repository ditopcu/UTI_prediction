# -*- coding: utf-8 -*-
"""
Renumber ALL references to a single sequential list REF_01, REF_02, ... in order of first
in-text citation, and rebuild the reference list to match. Merges the two legacy schemes
(REF1..REF13 and REF_N-01..REF_N-13) into one list.

Also bakes in two verified corrections:
  - Suresh: authors -> "Suresh K, Mohan D, Ananthanarayanan R, Niyas VKM" (one person, was mis-split)
  - Vickers: append doi:10.1177/0272989X06295361

Leaves bare "[REF]" / "(REF)" placeholders untouched (unresolved on purpose).
Re-runnable: recognizes REF_N-\\d+, REF_\\d+, REF\\d+ tag forms.

Target: paper/final versions/Draft DIT 2026.08.02  to Emilio.docx  (edit FILE below if renamed)
"""
import os, re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

BASE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
FILE = os.path.join(BASE, "paper", "final versions", "Draft DIT 2026.08.02  to Emilio.docx")

TAG_RE = re.compile(r'(REF_N-\d+|REF_\d+|REF\d+)(?![\d-])')

d = Document(FILE)
paras = d.paragraphs
ref_hi = next(i for i, p in enumerate(paras) if p.text.strip() == "REFERENCES")
ack_hi = next((i for i, p in enumerate(paras) if p.text.strip() == "ACKNOWLEDGEMENTS"), len(paras))

# ---- 1. first-citation order from in-text (paragraphs before REFERENCES) ----
order = []
for p in paras[:ref_hi]:
    for m in TAG_RE.finditer(p.text):
        t = m.group(1)
        if t not in order:
            order.append(t)
new_of = {t: "REF_%02d" % (i + 1) for i, t in enumerate(order)}
print("Unique in-text tags:", len(order))

# ---- 2. parse reference-list entries (between REFERENCES and ACKNOWLEDGEMENTS) ----
entry_re = re.compile(r'^\[?(REF_N-\d+|REF_\d+|REF\d+)\]?\.?\s*(.*)$', re.S)
entries = {}
for p in paras[ref_hi + 1:ack_hi]:
    t = p.text.strip()
    if not t:
        continue
    m = entry_re.match(t)
    if m:
        entries[m.group(1)] = m.group(2).strip()
print("Reference-list entries parsed:", len(entries))

# corrections
if "REF_N-13" in entries:
    entries["REF_N-13"] = entries["REF_N-13"].replace(
        "Suresh K, Mohan D, Ananthanarayanan R, Kandy V, Niyas M",
        "Suresh K, Mohan D, Ananthanarayanan R, Niyas VKM")
if "REF_N-08" in entries and "0272989X06295361" not in entries["REF_N-08"]:
    entries["REF_N-08"] = entries["REF_N-08"].rstrip()
    if not entries["REF_N-08"].endswith("."):
        entries["REF_N-08"] += "."
    entries["REF_N-08"] += " doi:10.1177/0272989X06295361."

# sanity: every cited tag has an entry
missing = [t for t in order if t not in entries]
if missing:
    raise SystemExit("Cited tags without a list entry: %s" % missing)

# ---- 3. replace in-text tags (single pass, collision-safe via lookup) ----
def repl(m):
    return new_of.get(m.group(1), m.group(0))
n_repl = 0
for p in paras[:ref_hi]:
    if TAG_RE.search(p.text):
        new_text = TAG_RE.sub(repl, p.text)
        if new_text != p.text:
            if p.runs:
                p.runs[0].text = new_text
                for r in p.runs[1:]:
                    r._r.getparent().remove(r._r)
            else:
                p.add_run(new_text)
            n_repl += 1
print("In-text paragraphs updated:", n_repl)

# ---- 4. rebuild reference list: remove old entry paragraphs + label, insert new ordered ----
ref_p = paras[ref_hi]
ack_p = paras[ack_hi] if ack_hi < len(paras) else None
# collect elements strictly between REFERENCES and ACKNOWLEDGEMENTS
el = ref_p._p.getnext()
to_remove = []
while el is not None and (ack_p is None or el is not ack_p._p):
    if el.tag == qn('w:p'):
        to_remove.append(el)
    el = el.getnext()
ref_style = None
for el in to_remove:
    # capture a style from an existing entry
    if ref_style is None:
        try: ref_style = Paragraph(el, ref_p._parent).style.name
        except Exception: pass
    el.getparent().remove(el)
ref_style = ref_style or "Normal"

def insert_after(anchor_p, text):
    np_ = OxmlElement('w:p'); anchor_p._p.addnext(np_)
    par = Paragraph(np_, anchor_p._parent)
    try: par.style = d.styles[ref_style]
    except KeyError: pass
    par.add_run(text)
    return par

prev = ref_p
for i, t in enumerate(order, 1):
    prev = insert_after(prev, "REF_%02d. %s" % (i, entries[t]))

d.save(FILE)
print("Saved. New list length:", len(order))
